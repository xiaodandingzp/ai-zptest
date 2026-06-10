"""知识库处理模块 - 支持多知识库"""

import os
import json
import shutil
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from .config import ConfigManager
from .vector_store import VectorStore, is_vectorstore_available

# 尝试导入 docx 支持
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class KnowledgeBaseManager:
    """多知识库管理类"""
    
    def __init__(self):
        self.config_manager = ConfigManager()
        self.data_dir = Path.home() / '.zpp' / 'knowledge_bases'
        self.config_file = self.data_dir / 'config.json'
        self.ensure_data_dir()
    
    def ensure_data_dir(self):
        """确保数据目录存在"""
        self.data_dir.mkdir(parents=True, exist_ok=True)
        if not self.config_file.exists():
            self.save_config({
                "bases": {},
                "selected_bases": []
            })
    
    def get_config(self) -> dict:
        """获取知识库配置"""
        try:
            with open(self.config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {"bases": {}, "selected_bases": []}
    
    def save_config(self, config: dict):
        """保存知识库配置"""
        with open(self.config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
    
    def create_knowledge_base(self, name: str, description: str = "") -> dict:
        """
        创建新知识库
        
        Args:
            name: 知识库名称
            description: 知识库描述
            
        Returns:
            创建结果
        """
        if not name or not name.strip():
            return {"success": False, "error": "知识库名称不能为空"}
        
        name = name.strip()
        
        # 生成唯一 ID
        import time
        base_id = f"kb_{int(time.time() * 1000)}"
        
        # 创建知识库目录
        base_dir = self.data_dir / base_id
        base_dir.mkdir(parents=True, exist_ok=True)
        
        # 创建文件目录
        files_dir = base_dir / "files"
        files_dir.mkdir(exist_ok=True)
        
        # 保存到配置
        config = self.get_config()
        config["bases"][base_id] = {
            "id": base_id,
            "name": name,
            "description": description,
            "created_at": time.time(),
            "file_count": 0,
            "chunk_count": 0,
            "initialized": False
        }
        self.save_config(config)
        
        return {
            "success": True,
            "message": f"知识库 '{name}' 创建成功",
            "base": config["bases"][base_id]
        }
    
    def delete_knowledge_base(self, base_id: str) -> dict:
        """删除知识库"""
        config = self.get_config()
        
        if base_id not in config["bases"]:
            return {"success": False, "error": "知识库不存在"}
        
        # 删除目录
        base_dir = self.data_dir / base_id
        if base_dir.exists():
            shutil.rmtree(base_dir)
        
        # 从配置中移除
        del config["bases"][base_id]
        
        # 从选中列表中移除
        if base_id in config.get("selected_bases", []):
            config["selected_bases"].remove(base_id)
        
        self.save_config(config)
        
        return {"success": True, "message": "知识库已删除"}
    
    def list_knowledge_bases(self) -> List[dict]:
        """列出所有知识库"""
        config = self.get_config()
        bases = []
        
        for base_id, base_info in config.get("bases", {}).items():
            base_info["selected"] = base_id in config.get("selected_bases", [])
            bases.append(base_info)
        
        # 按创建时间排序
        bases.sort(key=lambda x: x.get("created_at", 0), reverse=True)
        return bases
    
    def select_knowledge_base(self, base_id: str, selected: bool = True) -> dict:
        """选择/取消选择知识库"""
        config = self.get_config()
        
        if base_id not in config["bases"]:
            return {"success": False, "error": "知识库不存在"}
        
        if "selected_bases" not in config:
            config["selected_bases"] = []
        
        if selected:
            if base_id not in config["selected_bases"]:
                config["selected_bases"].append(base_id)
        else:
            if base_id in config["selected_bases"]:
                config["selected_bases"].remove(base_id)
        
        self.save_config(config)
        
        return {"success": True, "message": "已更新选择状态"}
    
    def get_selected_bases(self) -> List[str]:
        """获取选中的知识库 ID 列表"""
        config = self.get_config()
        return config.get("selected_bases", [])
    
    def get_knowledge_base_dir(self, base_id: str) -> Path:
        """获取知识库目录"""
        return self.data_dir / base_id / "files"
    
    def get_knowledge_base_info(self, base_id: str) -> dict:
        """获取知识库信息"""
        config = self.get_config()
        return config.get("bases", {}).get(base_id, {})
    
    def update_knowledge_base_info(self, base_id: str, info: dict):
        """更新知识库信息"""
        config = self.get_config()
        if base_id in config.get("bases", {}):
            config["bases"][base_id].update(info)
            self.save_config(config)


class KnowledgeManager:
    """单个知识库管理类"""
    
    def __init__(self, base_id: str = None):
        """
        初始化知识库管理器
        
        Args:
            base_id: 知识库 ID，如果为 None 则使用默认目录
        """
        self.config_manager = ConfigManager()
        self.base_manager = KnowledgeBaseManager()
        self.base_id = base_id
        
        if base_id:
            self.knowledge_dir = self.base_manager.get_knowledge_base_dir(base_id)
        else:
            # 默认使用 zpp 包下的 knowlage 目录
            package_dir = Path(__file__).parent
            self.knowledge_dir = package_dir / "knowlage"
        
        self.vector_store_dir = self._get_vector_store_dir()
        self.vector_store = None
    
    def _get_vector_store_dir(self) -> Path:
        """获取向量数据库目录"""
        if self.base_id:
            return Path.home() / '.zpp' / 'knowledge_bases' / self.base_id / 'vector_db'
        else:
            return Path.home() / '.zpp' / 'chroma_db'
    
    def scan_knowledge_files(self) -> List[Path]:
        """扫描知识库目录下的所有文件"""
        if not self.knowledge_dir.exists():
            return []
        
        supported_extensions = ['.txt', '.md', '.json', '.py', '.js', '.html', '.css', '.docx']
        files = []
        
        for ext in supported_extensions:
            files.extend(self.knowledge_dir.glob(f'*{ext}'))
        
        # 也支持子目录
        for ext in supported_extensions:
            files.extend(self.knowledge_dir.glob(f'**/*{ext}'))
        
        # 去重
        files = list(set(files))
        
        return files
    
    def read_file(self, file_path: Path) -> Tuple[str, dict]:
        """读取文件内容"""
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "extension": file_path.suffix,
            "base_id": self.base_id
        }
        
        try:
            if file_path.suffix.lower() == '.docx':
                if not DOCX_AVAILABLE:
                    print(f"跳过 {file_path}: 需要安装 python-docx")
                    return "", metadata
                return self._read_docx(file_path, metadata)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return content, metadata
        except Exception as e:
            print(f"读取文件失败 {file_path}: {e}")
            return "", metadata
    
    def _read_docx(self, file_path: Path, metadata: dict) -> Tuple[str, dict]:
        """读取 docx 文件内容"""
        try:
            doc = Document(str(file_path))
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            content = '\n\n'.join(paragraphs)
            return content, metadata
            
        except Exception as e:
            print(f"读取 docx 文件失败 {file_path}: {e}")
            return "", metadata
    
    def init_knowledge_base(self, verbose: bool = True) -> dict:
        """初始化知识库"""
        if not is_vectorstore_available():
            return {
                "success": False,
                "error": "向量数据库功能不可用，请安装依赖"
            }
        
        model_config = self.config_manager.get_current_model_config()
        if not model_config:
            return {"success": False, "error": "未配置模型"}
        
        api_key = model_config.get("api_key", "")
        api_base = model_config.get("api_base", "")
        
        if not api_key:
            return {"success": False, "error": "API Key 未配置"}
        
        files = self.scan_knowledge_files()
        if not files:
            return {"success": False, "error": f"知识库目录下没有找到支持的文件"}
        
        if verbose:
            print(f"找到 {len(files)} 个文件")
        
        texts = []
        metadatas = []
        
        for file_path in files:
            content, metadata = self.read_file(file_path)
            if content:
                texts.append(content)
                metadatas.append(metadata)
                if verbose:
                    print(f"已读取: {file_path.name}")
        
        if not texts:
            return {"success": False, "error": "没有成功读取任何文件内容"}
        
        try:
            self.vector_store = VectorStore(persist_directory=str(self.vector_store_dir))
            chunk_count = self.vector_store.create_from_texts(
                texts=texts,
                metadatas=metadatas,
                api_key=api_key,
                api_base=api_base
            )
            
            # 更新知识库信息
            if self.base_id:
                self.base_manager.update_knowledge_base_info(self.base_id, {
                    "file_count": len(files),
                    "chunk_count": chunk_count,
                    "initialized": True
                })
            
            return {
                "success": True,
                "file_count": len(files),
                "chunk_count": chunk_count
            }
            
        except Exception as e:
            return {"success": False, "error": f"创建向量数据库失败: {str(e)}"}
    
    def search(self, query: str, k: int = 4) -> List[dict]:
        """搜索相关知识"""
        if not is_vectorstore_available():
            return []
        
        try:
            model_config = self.config_manager.get_current_model_config()
            if not model_config:
                return []
            
            api_key = model_config.get("api_key", "")
            api_base = model_config.get("api_base", "")
            
            if not api_key:
                return []
            
            if self.vector_store is None:
                self.vector_store = VectorStore(persist_directory=str(self.vector_store_dir))
                self.vector_store.load_vectorstore(api_key, api_base)
            
            return self.vector_store.similarity_search(query, k=k)
            
        except Exception as e:
            print(f"搜索失败: {e}")
            return []
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['.txt', '.md', '.json', '.py', '.js', '.html', '.css', '.docx']
    
    def is_supported_file(self, filename: str) -> bool:
        """检查文件是否支持"""
        ext = Path(filename).suffix.lower()
        return ext in self.get_supported_extensions()
    
    def list_knowledge_files(self) -> List[dict]:
        """列出知识库目录下的所有文件信息"""
        if not self.knowledge_dir.exists():
            return []
        
        files = []
        for file_path in self.knowledge_dir.iterdir():
            if file_path.is_file():
                stat = file_path.stat()
                files.append({
                    "name": file_path.name,
                    "path": str(file_path),
                    "extension": file_path.suffix,
                    "size": stat.st_size,
                    "size_display": self._format_file_size(stat.st_size),
                    "supported": self.is_supported_file(file_path.name),
                    "modified": stat.st_mtime
                })
        
        files.sort(key=lambda x: x.get("modified", 0), reverse=True)
        return files
    
    def _format_file_size(self, size: int) -> str:
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"
    
    def add_file(self, filename: str, content: bytes) -> dict:
        """添加新文件到知识库"""
        if not filename:
            return {"success": False, "error": "文件名不能为空"}
        
        if not self.is_supported_file(filename):
            return {
                "success": False, 
                "error": f"不支持的文件格式"
            }
        
        file_path = self.knowledge_dir / filename
        if file_path.exists():
            return {"success": False, "error": "文件已存在"}
        
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            with open(file_path, 'wb') as f:
                f.write(content)
            
            return {
                "success": True, 
                "message": f"文件 {filename} 添加成功",
                "file": {"name": filename, "size": len(content)}
            }
        except Exception as e:
            return {"success": False, "error": f"写入文件失败: {str(e)}"}
    
    def delete_file(self, filename: str) -> dict:
        """删除知识库中的文件"""
        file_path = self.knowledge_dir / filename
        
        if not file_path.exists():
            return {"success": False, "error": "文件不存在"}
        
        try:
            file_path.resolve().relative_to(self.knowledge_dir.resolve())
        except ValueError:
            return {"success": False, "error": "非法文件路径"}
        
        try:
            file_path.unlink()
            return {"success": True, "message": f"文件 {filename} 已删除"}
        except Exception as e:
            return {"success": False, "error": f"删除文件失败: {str(e)}"}
    
    def clear_knowledge_base(self):
        """清空知识库向量数据"""
        if self.vector_store_dir.exists():
            shutil.rmtree(self.vector_store_dir)
        self.vector_store = None
        
        if self.base_id:
            self.base_manager.update_knowledge_base_info(self.base_id, {
                "initialized": False,
                "chunk_count": 0
            })


def search_selected_knowledge_bases(query: str, k: int = 4) -> List[dict]:
    """
    搜索所有选中的知识库
    
    Args:
        query: 查询文本
        k: 每个知识库返回的结果数量
        
    Returns:
        合并后的搜索结果
    """
    base_manager = KnowledgeBaseManager()
    selected_ids = base_manager.get_selected_bases()
    
    if not selected_ids:
        return []
    
    all_results = []
    
    for base_id in selected_ids:
        manager = KnowledgeManager(base_id)
        results = manager.search(query, k=k)
        
        # 添加来源信息
        base_info = base_manager.get_knowledge_base_info(base_id)
        for result in results:
            result["base_name"] = base_info.get("name", "未知知识库")
            result["base_id"] = base_id
        
        all_results.extend(results)
    
    # 按相关性排序（这里简化处理，实际应该按相似度分数排序）
    return all_results[:k * len(selected_ids)]


def get_context_from_selected_bases(query: str, k: int = 4) -> str:
    """从选中的知识库获取上下文"""
    results = search_selected_knowledge_bases(query, k)
    
    if not results:
        return ""
    
    context_parts = []
    for i, item in enumerate(results, 1):
        source = item.get("metadata", {}).get("filename", "未知来源")
        base_name = item.get("base_name", "未知知识库")
        content = item.get("content", "")
        context_parts.append(f"[参考资料 {i} - {base_name}/{source}]\n{content}\n")
    
    return "\n".join(context_parts)
